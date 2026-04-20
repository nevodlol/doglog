from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from typing import List, Optional
import shutil, os
from io import BytesIO
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from copy import deepcopy
from pathlib import Path
from database import SessionLocal
from models import Dog, GenderEnum
from schemas import DogResponse, DogUpdate

router = APIRouter(prefix="/dogs", tags=["dogs"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
TEMPLATE_PATH = Path(__file__).resolve().parents[2] / "акт взвешивания.docx"

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@router.get("/", response_model=List[DogResponse])
def get_dogs(db: Session = Depends(get_db)):
    return db.query(Dog).all()

@router.get("/weights-report/download")
def download_weights_report(db: Session = Depends(get_db)):
    dogs = db.query(Dog).order_by(Dog.id.asc()).all()
    if not TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="Шаблон акта не найден")

    doc = Document(str(TEMPLATE_PATH))
    placeholder_paragraphs = [p for p in doc.paragraphs if "{dogs.name}" in p.text]

    if not placeholder_paragraphs:
        raise HTTPException(status_code=500, detail="В шаблоне нет полей для собак")

    base_paragraph = placeholder_paragraphs[0]
    base_template = base_paragraph.text

    def render_dog_line(template: str, dog: Dog):
        gender_map = {
            "male": "кобель",
            "female": "сука",
        }
        raw_gender = dog.gender.value if dog.gender else ""
        gender_value = gender_map.get(raw_gender, raw_gender)
        weight_value = "" if dog.weight is None else f"{dog.weight:.3f}".replace(".", ",")
        rendered = (
            template
            .replace("{dogs.name}", dog.name or "")
            .replace("{dogs.breed}", dog.breed or "")
            .replace("{dogs.gender}", gender_value)
            .replace("{dogs.weight}", weight_value)
            .replace("{dogs.count}", str(len(dogs)))
            .replace("{date.today}", date.today().strftime("%d.%m.%Y"))
        )
        # В шаблоне может быть ручной перенос перед тире/весом.
        # Приводим строку к однострочному виду, чтобы вес шел сразу после вводных.
        rendered = rendered.replace("\r", " ").replace("\n", " ")
        # Запятую перед тире/весом заменяем на пробел: «…кобель –12,340», а не «…кобель,–12,340».
        rendered = (
            rendered.replace(", –", " –")
            .replace(", -", " -")
            .replace(",–", " –")
            .replace(",-", " -")
        )
        return rendered

    if dogs:
        base_paragraph.text = render_dog_line(base_template, dogs[0])
        base_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        anchor_paragraph = base_paragraph

        for idx in range(1, len(dogs)):
            new_p = deepcopy(base_paragraph._p)
            anchor_paragraph._p.addnext(new_p)
            new_paragraph = Paragraph(new_p, base_paragraph._parent)
            new_paragraph.text = render_dog_line(base_template, dogs[idx])
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            anchor_paragraph = new_paragraph
    else:
        base_paragraph.text = "Собаки в базе отсутствуют."
        base_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for paragraph in placeholder_paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)

    for paragraph in doc.paragraphs:
        if "{dogs.count}" in paragraph.text or "{date.today}" in paragraph.text:
            paragraph.text = (
                paragraph.text
                .replace("{dogs.count}", str(len(dogs)))
                .replace("{date.today}", date.today().strftime("%d.%m.%Y"))
            )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=akt_vzveshivaniya.docx"
        },
    )

@router.post("/", response_model=DogResponse)
async def create_dog(
    name: str = Form(...),
    gender: GenderEnum = Form(...),
    birthdate: str = Form(...),
    breed: str = Form(...),
    color: str = Form(...),
    chip: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db)
):

    file_path = None
    if file:
        file_path = f"{UPLOAD_DIR}/{file.filename}"
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

    new_dog = Dog(
        name=name,
        gender=gender,
        birthdate=birthdate,
        breed=breed,
        color=color,
        chip=chip,
        photo=file_path
    )

    db.add(new_dog)
    db.commit()
    db.refresh(new_dog)

    return new_dog

@router.delete("/{dog_id}")
def delete_dog(dog_id: int, db: Session = Depends(get_db)):
    dog = db.query(Dog).filter(Dog.id == dog_id).first()

    if not dog:
        raise HTTPException(status_code=404, detail="Собака не найдена")

    if dog.photo:
        try:
            os.remove(dog.photo)
        except FileNotFoundError:
            pass

    db.delete(dog)
    db.commit()

    return {"message": "Собака удалена", "id": dog_id}

@router.patch("/{dog_id}", response_model=DogResponse)
def update_dog(dog_id: int, data: DogUpdate, db: Session = Depends(get_db)):
    dog = db.query(Dog).filter(Dog.id == dog_id).first()

    if not dog:
        raise HTTPException(status_code=404, detail="Собака не найдена")

    update_data = data.model_dump(exclude_unset=True)

    for key, value in update_data.items():
        setattr(dog, key, value)

    db.commit()
    db.refresh(dog)

    return dog

@router.put("/{dog_id}", response_model=DogResponse)
def replace_dog(dog_id: int, data: DogUpdate, db: Session = Depends(get_db)):
    dog = db.query(Dog).filter(Dog.id == dog_id).first()

    if not dog:
        raise HTTPException(status_code=404, detail="Собака не найдена")

    update_data = data.model_dump(exclude_unset=True)

    for key, value in update_data.items():
        setattr(dog, key, value)

    db.commit()
    db.refresh(dog)

    return dog